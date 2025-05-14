import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from clothing_factory import ClothingFactory, Season, ClothingType, Outerwear, Underwear, PantsShorts, Top, MensSet, WomensSet
from PIL import Image, ImageTk
import os
import psycopg2
from datetime import datetime
from psycopg2 import Error
import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
import random
import string

class АссортиментМагазинаGUI:
    def __init__(self, parent, cursor, conn):
        self.window = tk.Toplevel(parent)
        self.window.title("Ассортимент магазина")
        self.window.geometry("1200x800")

        self.cursor = cursor
        self.conn = conn

        # Основной контейнер
        main_frame = ttk.Frame(self.window)
        main_frame.pack(expand=True, fill='both', padx=20, pady=20)

        # Фильтры
        frame_фильтры = ttk.LabelFrame(main_frame, text="Фильтры", padding=10)
        frame_фильтры.pack(fill='x', padx=5, pady=5)

        # Поиск по штрих-коду
        frame_поиск = ttk.Frame(frame_фильтры)
        frame_поиск.pack(fill='x', pady=5)
        ttk.Label(frame_поиск, text="Поиск по штрих-коду:").pack(side='left', padx=5)
        self.поиск_штрих_код = ttk.Entry(frame_поиск, width=20)
        self.поиск_штрих_код.pack(side='left', padx=5)
        self.поиск_штрих_код.bind('<KeyRelease>', self.поиск_по_штрих_коду_в_реальном_времени)
        ttk.Button(frame_поиск, text="Сбросить", command=self.сбросить_поиск).pack(side='left', padx=5)

        # Фильтр по типу
        ttk.Label(frame_фильтры, text="Тип:").pack(side='left', padx=5)
        self.фильтр_тип = ttk.Combobox(frame_фильтры, values=[
            "Все",
            "Верхняя одежда",
            "Нижнее белье",
            "Штаны/Шорты",
            "Верхняя нательная одежда",
            "Комплекты муж",
            "Комплекты жен"
        ], width=25)
        self.фильтр_тип.pack(side='left', padx=5)
        self.фильтр_тип.set("Все")

        # Фильтр по размеру
        ttk.Label(frame_фильтры, text="Размер:").pack(side='left', padx=5)
        self.фильтр_размер = ttk.Combobox(frame_фильтры, values=[
            "Все",
            "XS", "S", "M", "L", "XL", "XXL"
        ], width=15)
        self.фильтр_размер.pack(side='left', padx=5)
        self.фильтр_размер.set("Все")

        # Фильтр по цвету
        ttk.Label(frame_фильтры, text="Цвет:").pack(side='left', padx=5)
        self.фильтр_цвет = ttk.Combobox(frame_фильтры, values=[
            "Все",
            "Черный", "Белый", "Красный", "Синий", "Зеленый", "Желтый", "Серый", "Коричневый"
        ], width=20)
        self.фильтр_цвет.pack(side='left', padx=5)
        self.фильтр_цвет.set("Все")

        # Кнопка применения фильтров
        ttk.Button(frame_фильтры, text="Применить фильтры",
                  command=self.обновить_список_товаров).pack(side='left', padx=5)

        # Список товаров
        frame_список = ttk.LabelFrame(main_frame, text="Ассортимент магазина", padding=10)
        frame_список.pack(fill='both', expand=True, padx=5, pady=5)

        self.список_товаров = ttk.Treeview(frame_список, columns=("ID", "Штрих-код", "Название", "Тип", "Размер", "Цвет", "Количество", "Цена"), show="headings", height=15)
        self.список_товаров.pack(expand=True, fill='both')

        # Увеличиваем ширину колонок
        self.список_товаров.column("ID", width=50)
        self.список_товаров.column("Штрих-код", width=120)
        self.список_товаров.column("Название", width=200)
        self.список_товаров.column("Тип", width=150)
        self.список_товаров.column("Размер", width=100)
        self.список_товаров.column("Цвет", width=100)
        self.список_товаров.column("Количество", width=100)
        self.список_товаров.column("Цена", width=100)

        for col in ("ID", "Штрих-код", "Название", "Тип", "Размер", "Цвет", "Количество", "Цена"):
            self.список_товаров.heading(col, text=col)

        # Кнопки управления
        frame_кнопки = ttk.Frame(frame_список)
        frame_кнопки.pack(fill='x', pady=10)

        # Создаем фрейм для левой группы кнопок
        frame_кнопки_лево = ttk.Frame(frame_кнопки)
        frame_кнопки_лево.pack(side='left', padx=5)
        ttk.Button(frame_кнопки_лево, text="Продать", command=self.продать_товар, width=20).pack(side='left', padx=5)
        ttk.Button(frame_кнопки_лево, text="Информация", command=self.показать_информацию, width=20).pack(side='left', padx=5)

        # Создаем фрейм для правой группы кнопок
        frame_кнопки_право = ttk.Frame(frame_кнопки)
        frame_кнопки_право.pack(side='right', padx=5)
        ttk.Button(frame_кнопки_право, text="Экспорт в Word", command=self.экспорт_в_word, width=20).pack(side='left', padx=5)
        ttk.Button(frame_кнопки_право, text="Экспорт полной документации", command=self.экспорт_документации, width=25).pack(side='left', padx=5)

        # Метка для общей стоимости
        self.метка_стоимости = ttk.Label(main_frame, text="", font=('Arial', 14, 'bold'))
        self.метка_стоимости.pack(pady=10)

        # Обновляем список при создании интерфейса
        self.обновить_список_товаров()

    def поиск_по_штрих_коду_в_реальном_времени(self, event=None):
        штрих_код = self.поиск_штрих_код.get().strip()

        # Очищаем список
        for item in self.список_товаров.get_children():
            self.список_товаров.delete(item)

        if not штрих_код:
            self.обновить_список_товаров()
            return

        # Ищем товары, у которых штрих-код начинается с введенных цифр
        self.cursor.execute("""
            SELECT id, barcode, name, type, size, color, quantity, price
            FROM clothing
            WHERE material = 'Ассортимент магазина' AND barcode LIKE %s
            ORDER BY barcode
        """, (f"{штрих_код}%",))

        total = 0
        for товар in self.cursor.fetchall():
            self.список_товаров.insert("", "end", values=(
                товар[0],  # ID
                товар[1],  # Штрих-код
                товар[2],  # Название
                товар[3],  # Тип
                товар[4],  # Размер
                товар[5],  # Цвет
                товар[6],  # Количество
                f"{товар[7]} руб."  # Цена
            ))
            try:
                qty = int(товар[6]) if товар[6] is not None else 1
                price = float(товар[7]) if товар[7] is not None else 0
                total += qty * price
            except Exception:
                pass

        self.метка_стоимости.configure(text=f"Общая стоимость ассортимента: {total:.2f} руб.")

    def сбросить_поиск(self):
        self.поиск_штрих_код.delete(0, tk.END)
        self.обновить_список_товаров()

    def обновить_список_товаров(self):
        # Очищаем список
        for item in self.список_товаров.get_children():
            self.список_товаров.delete(item)

        # Получаем значения фильтров
        выбранный_тип = self.фильтр_тип.get()
        выбранный_размер = self.фильтр_размер.get()
        выбранный_цвет = self.фильтр_цвет.get()

        # Формируем SQL-запрос
        query = """
            SELECT id, barcode, name, type, size, color, quantity, price
            FROM clothing
            WHERE material = 'Ассортимент магазина'
        """
        params = []

        # Добавляем условия фильтрации
        if выбранный_тип and выбранный_тип != "Все":
            query += " AND type = %s"
            params.append(выбранный_тип)
        if выбранный_размер and выбранный_размер != "Все":
            query += " AND size = %s"
            params.append(выбранный_размер)
        if выбранный_цвет and выбранный_цвет != "Все":
            query += " AND color = %s"
            params.append(выбранный_цвет)

        query += " ORDER BY created_at DESC"

        # Выполняем запрос
        self.cursor.execute(query, params)

        total = 0
        for товар in self.cursor.fetchall():
            self.список_товаров.insert("", "end", values=(
                товар[0],  # ID
                товар[1],  # Штрих-код
                товар[2],  # Название
                товар[3],  # Тип
                товар[4],  # Размер
                товар[5],  # Цвет
                товар[6],  # Количество
                f"{товар[7]} руб."  # Цена
            ))
            try:
                qty = int(товар[6]) if товар[6] is not None else 1
                price = float(товар[7]) if товар[7] is not None else 0
                total += qty * price
            except Exception:
                pass

        self.метка_стоимости.configure(text=f"Общая стоимость ассортимента: {total:.2f} руб.")

    def продать_товар(self):
        selected_items = self.список_товаров.selection()
        if not selected_items:
            messagebox.showwarning("Предупреждение", "Выберите товар для продажи")
            return

        # Создаем окно продажи
        окно_продажи = tk.Toplevel(self.window)
        окно_продажи.title("Продажа товара")
        окно_продажи.geometry("500x400")
        окно_продажи.resizable(False, False)

        # Получаем данные о выбранном товаре
        item = self.список_товаров.item(selected_items[0])
        values = item['values']

        # Отображаем информацию о товаре
        frame_информация = ttk.LabelFrame(окно_продажи, text="Информация о товаре", padding=15)
        frame_информация.pack(fill='x', padx=15, pady=10)

        style = ttk.Style()
        style.configure('Info.TLabel', font=('Arial', 12))

        ttk.Label(frame_информация, text=f"Название: {values[2]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Тип: {values[3]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Размер: {values[4]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Цвет: {values[5]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Доступное количество: {values[6]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Цена: {values[7]}", style='Info.TLabel').pack(anchor='w', pady=3)

        # Поле для ввода количества
        frame_количество = ttk.LabelFrame(окно_продажи, text="Количество для продажи", padding=15)
        frame_количество.pack(fill='x', padx=15, pady=10)

        количество = ttk.Entry(frame_количество, width=15, font=('Arial', 12))
        количество.pack(pady=10)
        количество.insert(0, "1")

        def подтвердить_продажу():
            try:
                кол = int(количество.get())
                if кол <= 0:
                    raise ValueError("Количество должно быть положительным числом")

                доступное_количество = int(values[6])
                if кол > доступное_количество:
                    raise ValueError(f"Недостаточно товара. Доступно: {доступное_количество}")

                # Получаем ID товара из базы данных
                self.cursor.execute("""
                    SELECT id, price FROM clothing
                    WHERE name = %s AND type = %s AND size = %s AND color = %s AND material = 'Ассортимент магазина'
                    LIMIT 1
                """, (values[2], values[3], values[4], values[5]))
                row = self.cursor.fetchone()
                if not row:
                    raise ValueError("Товар не найден в базе данных")

                clothing_id, price = row
                total_price = price * кол

                # Добавляем запись о продаже
                self.cursor.execute("""
                    INSERT INTO sales (clothing_id, type, size, color, price, quantity, total_price)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (clothing_id, values[3], values[4], values[5], price, кол, total_price))

                # Уменьшаем количество товара
                self.cursor.execute("""
                    UPDATE clothing
                    SET quantity = quantity - %s
                    WHERE id = %s
                """, (кол, clothing_id))

                self.conn.commit()

                # Обновляем список товаров
                self.обновить_список_товаров()

                messagebox.showinfo("Успех", f"Товар успешно продан!\nКоличество: {кол}\nОбщая сумма: {total_price} руб.")
                окно_продажи.destroy()

            except ValueError as e:
                messagebox.showerror("Ошибка", str(e))
            except Error as e:
                self.conn.rollback()
                messagebox.showerror("Ошибка базы данных", str(e))

        # Кнопки
        frame_кнопки = ttk.Frame(окно_продажи)
        frame_кнопки.pack(pady=15)
        ttk.Button(frame_кнопки, text="Подтвердить продажу", command=подтвердить_продажу, width=20).pack(side='left', padx=10)
        ttk.Button(frame_кнопки, text="Отмена", command=окно_продажи.destroy, width=20).pack(side='left', padx=10)

        # Делаем окно модальным
        окно_продажи.transient(self.window)
        окно_продажи.grab_set()
        self.window.wait_window(окно_продажи)

    def показать_информацию(self):
        selected_items = self.список_товаров.selection()
        if not selected_items:
            messagebox.showwarning("Предупреждение", "Выберите товар для просмотра информации")
            return

        # Получаем данные о выбранном товаре
        item = self.список_товаров.item(selected_items[0])
        values = item['values']

        # Создаем окно с информацией
        info_window = tk.Toplevel(self.window)
        info_window.title("Информация о товаре")
        info_window.geometry("400x500")
        info_window.resizable(False, False)

        # Создаем фрейм для информации
        frame_информация = ttk.Frame(info_window, padding=20)
        frame_информация.pack(fill='both', expand=True)

        # Стиль для меток
        style = ttk.Style()
        style.configure('Info.TLabel', font=('Arial', 12))

        ttk.Label(frame_информация, text=f"ID: {values[0]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Штрих-код: {values[1]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Название: {values[2]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Тип: {values[3]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Размер: {values[4]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Цвет: {values[5]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Доступное количество: {values[6]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Цена: {values[7]}", style='Info.TLabel').pack(anchor='w', pady=3)

    def экспорт_в_word(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if not file_path:
            return

        doc = Document()
        doc.add_heading("Ассортимент магазина", 0)

        table = doc.add_table(rows=1, cols=8)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Штрих-код'
        hdr_cells[2].text = 'Название'
        hdr_cells[3].text = 'Тип'
        hdr_cells[4].text = 'Размер'
        hdr_cells[5].text = 'Цвет'
        hdr_cells[6].text = 'Количество'
        hdr_cells[7].text = 'Цена'

        self.cursor.execute("""
            SELECT id, barcode, name, type, size, color, quantity, price
            FROM clothing
            WHERE material = 'Ассортимент магазина'
            ORDER BY created_at DESC
        """)
        for row in self.cursor.fetchall():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        doc.save(file_path)
        messagebox.showinfo("Экспорт завершен", f"Данные успешно экспортированы в {file_path}")

    def экспорт_документации(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if not file_path:
            return

        doc = Document()
        doc.add_heading("Документация по ассортименту магазина", 0)

        # 1. Ассортимент магазина
        doc.add_heading("Ассортимент магазина", level=1)
        table_ассортимент = doc.add_table(rows=1, cols=8)
        hdr_cells = table_ассортимент.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Штрих-код'
        hdr_cells[2].text = 'Название'
        hdr_cells[3].text = 'Тип'
        hdr_cells[4].text = 'Размер'
        hdr_cells[5].text = 'Цвет'
        hdr_cells[6].text = 'Количество'
        hdr_cells[7].text = 'Цена'

        self.cursor.execute("""
            SELECT id, barcode, name, type, size, color, quantity, price
            FROM clothing
            WHERE material = 'Ассортимент магазина'
            ORDER BY type, size
        """)
        total_ассортимент = 0
        for row in self.cursor.fetchall():
            row_cells = table_ассортимент.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
            total_ассортимент += row[6] * row[7] if row[6] and row[7] else 0

        doc.add_paragraph(f"Общая стоимость ассортимента: {total_ассортимент:.2f} руб.")

        # 2. Продажи
        doc.add_heading("История продаж", level=1)
        table_продажи = doc.add_table(rows=1, cols=8)
        hdr_cells = table_продажи.rows[0].cells
        hdr_cells[0].text = 'Название'
        hdr_cells[1].text = 'Тип'
        hdr_cells[2].text = 'Размер'
        hdr_cells[3].text = 'Цвет'
        hdr_cells[4].text = 'Количество'
        hdr_cells[5].text = 'Цена за ед.'
        hdr_cells[6].text = 'Общая сумма'
        hdr_cells[7].text = 'Дата продажи'

        self.cursor.execute("""
            SELECT c.name, s.type, s.size, s.color, s.quantity, s.price, s.total_price, s.sale_date
            FROM sales s
            JOIN clothing c ON s.clothing_id = c.id
            ORDER BY s.sale_date DESC
        """)
        total_продажи = 0
        for row in self.cursor.fetchall():
            row_cells = table_продажи.add_row().cells
            for i, value in enumerate(row):
                if i == 7:  # Дата продажи
                    row_cells[i].text = value.strftime("%Y-%m-%d %H:%M")
                else:
                    row_cells[i].text = str(value)
            total_продажи += row[6] if row[6] else 0

        doc.add_paragraph(f"Общая сумма продаж: {total_продажи:.2f} руб.")

        # Добавляем итоговую информацию
        doc.add_heading("Итоговая информация", level=1)
        doc.add_paragraph(f"Общая стоимость ассортимента: {total_ассортимент:.2f} руб.")
        doc.add_paragraph(f"Общая сумма продаж: {total_продажи:.2f} руб.")

        # Сохраняем документ
        doc.save(file_path)
        messagebox.showinfo("Экспорт завершен", f"Документация успешно экспортирована в {file_path}")

class ТрикотажнаяФабрикаGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Трикотажная фабрика")
        self.root.geometry("1400x900")

        # Подключение к базе данных
        try:
            self.conn = psycopg2.connect(
                database=os.getenv('DB_NAME', 'clothing_factory'),
                user=os.getenv('DB_USER', 'postgres'),
                password=os.getenv('DB_PASSWORD', 'Ostap_628'),
                host=os.getenv('DB_HOST', 'localhost'),
                port=os.getenv('DB_PORT', '5432')
            )
            self.cursor = self.conn.cursor()
            self.создать_таблицы()
        except Error as e:
            messagebox.showerror("Ошибка подключения к БД", f"Не удалось подключиться к базе данных: {e}")
            self.root.destroy()
            return

        # Устанавливаем стиль
        self.style = ttk.Style()
        self.style.configure('TFrame', background='#f0f0f0')
        self.style.configure('TLabel', background='#f0f0f0', font=('Arial', 12))
        self.style.configure('TButton', font=('Arial', 12))
        self.style.configure('TCombobox', font=('Arial', 12))
        self.style.configure('TEntry', font=('Arial', 12))
        self.style.configure('Treeview', font=('Arial', 12))
        self.style.configure('Treeview.Heading', font=('Arial', 12, 'bold'))

        self.фабрика = ClothingFactory()

        # Создаем вкладки
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(expand=True, fill='both', padx=10, pady=5)

        # Вкладка добавления одежды
        self.frame_добавление = ttk.Frame(self.notebook)
        self.notebook.add(self.frame_добавление, text="Добавить одежду")

        # Вкладка склада
        self.frame_склад = ttk.Frame(self.notebook)
        self.notebook.add(self.frame_склад, text="Склад")

        # Вкладка поставщиков
        self.frame_поставщики = ttk.Frame(self.notebook)
        self.notebook.add(self.frame_поставщики, text="Поставщики")

        # Вкладка фурнитуры
        self.frame_фурнитура = ttk.Frame(self.notebook)
        self.notebook.add(self.frame_фурнитура, text="Фурнитура")

        # Вкладка готовых изделий
        self.frame_готовые_изделия = ttk.Frame(self.notebook)
        self.notebook.add(self.frame_готовые_изделия, text="Ассортимент магазина")

        # Вкладка склада фурнитуры
        self.frame_склад_фурнитуры = ttk.Frame(self.notebook)
        self.notebook.add(self.frame_склад_фурнитуры, text="Склад фурнитуры")

        # Вкладка продаж
        self.frame_продажи = ttk.Frame(self.notebook)
        self.notebook.add(self.frame_продажи, text="Продажи")

        self.создать_интерфейс_добавления()
        self.создать_интерфейс_склада()
        self.создать_интерфейс_поставщиков()
        self.создать_интерфейс_фурнитуры()
        self.создать_интерфейс_готовых_изделий()
        self.создать_интерфейс_склада_фурнитуры()
        self.создать_интерфейс_продаж()

        # Хранилище для изображений
        self.изображения = {}

        # Добавляем кнопку для тестовых данных
        frame_тестовые_данные = ttk.Frame(self.root)
        frame_тестовые_данные.pack(fill='x', padx=10, pady=5)
        ttk.Button(frame_тестовые_данные, text="Добавить тестовые данные",
                  command=self.добавить_тестовые_данные).pack(side='right')

    def создать_таблицы(self):
        # Создаем таблицы, если они не существуют
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS clothing (
                id SERIAL PRIMARY KEY,
                name VARCHAR(100) NOT NULL,
                type VARCHAR(50) NOT NULL,
                size VARCHAR(10) NOT NULL,
                color VARCHAR(50) NOT NULL,
                material VARCHAR(50) NOT NULL,
                price DECIMAL(10,2) NOT NULL,
                image_path TEXT,
                quantity INTEGER DEFAULT 1,
                barcode VARCHAR(13) UNIQUE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # Добавляем столбец barcode, если его нет
        try:
            self.cursor.execute("""
                ALTER TABLE clothing
                ADD COLUMN IF NOT EXISTS barcode VARCHAR(13) UNIQUE
            """)
            self.conn.commit()
        except Error as e:
            self.conn.rollback()
            print(f"Ошибка при добавлении столбца barcode: {e}")

        # Обновляем существующие записи, добавляя штрих-коды
        try:
            self.cursor.execute("SELECT id FROM clothing WHERE barcode IS NULL")
            records = self.cursor.fetchall()
            for record in records:
                barcode = self.сгенерировать_штрих_код()
                self.cursor.execute("""
                    UPDATE clothing
                    SET barcode = %s
                    WHERE id = %s
                """, (barcode, record[0]))
            self.conn.commit()
        except Error as e:
            self.conn.rollback()
            print(f"Ошибка при обновлении штрих-кодов: {e}")

        # Создаем таблицу для связи одежды и фурнитуры
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS clothing_fittings (
                id SERIAL PRIMARY KEY,
                clothing_id INTEGER REFERENCES clothing(id),
                fitting_type VARCHAR(50) NOT NULL,
                quantity INTEGER NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # Создаем таблицу для связи одежды и трикотажных полотен
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS clothing_knit_fabrics (
                id SERIAL PRIMARY KEY,
                clothing_id INTEGER REFERENCES clothing(id),
                fabric_type VARCHAR(50) NOT NULL,
                quantity INTEGER NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # Добавляем столбцы season и quantity, если их нет
        try:
            self.cursor.execute("""
                ALTER TABLE clothing
                ADD COLUMN IF NOT EXISTS season VARCHAR(20),
                ADD COLUMN IF NOT EXISTS quantity INTEGER DEFAULT 1
            """)
            self.conn.commit()
        except Error as e:
            self.conn.rollback()
            print(f"Ошибка при добавлении столбцов: {e}")

        # Создаем таблицу поставщиков
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS suppliers (
                id SERIAL PRIMARY KEY,
                company_name VARCHAR(100) NOT NULL,
                contact_person VARCHAR(100) NOT NULL,
                phone VARCHAR(20) NOT NULL,
                email VARCHAR(100) NOT NULL,
                address TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # Создаем таблицу заказчиков
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS customers (
                id SERIAL PRIMARY KEY,
                company_name VARCHAR(100) NOT NULL,
                contact_person VARCHAR(100) NOT NULL,
                phone VARCHAR(20) NOT NULL,
                email VARCHAR(100) NOT NULL,
                address TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # Создаем таблицу для хранения поступившей фурнитуры и трикотажа
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS received_fittings (
                id SERIAL PRIMARY KEY,
                order_id INTEGER,
                name VARCHAR(100),
                type VARCHAR(50) NOT NULL,
                quantity INTEGER NOT NULL,
                price DECIMAL(10,2) NOT NULL,
                status VARCHAR(20) DEFAULT 'В обработке',
                received_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # Снимаем ограничение NOT NULL с name, если оно есть
        try:
            self.cursor.execute("""
                ALTER TABLE received_fittings
                ALTER COLUMN name DROP NOT NULL
            """)
            self.conn.commit()
        except Error as e:
            self.conn.rollback()
            print(f"Ошибка при снятии ограничения NOT NULL с name: {e}")

        # Добавляем столбец status, если его нет
        try:
            self.cursor.execute("""
                ALTER TABLE received_fittings
                ADD COLUMN IF NOT EXISTS status VARCHAR(20) DEFAULT 'В обработке'
            """)
            self.conn.commit()
        except Error as e:
            self.conn.rollback()
            print(f"Ошибка при добавлении столбца status: {e}")

        # Создаем таблицу для продаж
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS sales (
                id SERIAL PRIMARY KEY,
                clothing_id INTEGER REFERENCES clothing(id),
                type VARCHAR(50) NOT NULL,
                size VARCHAR(10) NOT NULL,
                color VARCHAR(50) NOT NULL,
                price DECIMAL(10,2) NOT NULL,
                quantity INTEGER NOT NULL,
                total_price DECIMAL(10,2) NOT NULL,
                sale_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        self.conn.commit()

    def __del__(self):
        # Закрываем соединение с базой данных при уничтожении объекта
        if hasattr(self, 'conn'):
            self.conn.close()

    def создать_интерфейс_добавления(self):
        # Основной контейнер
        main_frame = ttk.Frame(self.frame_добавление)
        main_frame.pack(expand=True, fill='both', padx=10, pady=10)

        # Левая часть - форма
        frame_форма = ttk.LabelFrame(main_frame, text="Информация об одежде", padding=5)
        frame_форма.pack(side='left', fill='both', expand=True, padx=5, pady=5)

        # Тип одежды
        ttk.Label(frame_форма, text="Тип одежды:").grid(row=0, column=0, padx=5, pady=2, sticky='w')
        self.тип_одежды = ttk.Combobox(frame_форма, values=[
            "Верхняя одежда",
            "Нижнее белье",
            "Штаны/Шорты",
            "Верхняя нательная одежда",
            "Комплекты муж",
            "Комплекты жен"
        ], width=25)
        self.тип_одежды.grid(row=0, column=1, padx=5, pady=2)

        # Название
        ttk.Label(frame_форма, text="Название:").grid(row=1, column=0, padx=5, pady=2, sticky='w')
        self.название = ttk.Entry(frame_форма, width=25)
        self.название.grid(row=1, column=1, padx=5, pady=2)

        # Размер
        ttk.Label(frame_форма, text="Размер:").grid(row=2, column=0, padx=5, pady=2, sticky='w')
        self.размер = ttk.Combobox(frame_форма, values=["XS", "S", "M", "L", "XL", "XXL"], width=25)
        self.размер.grid(row=2, column=1, padx=5, pady=2)

        # Цвет
        ttk.Label(frame_форма, text="Цвет:").grid(row=3, column=0, padx=5, pady=2, sticky='w')
        self.цвет = ttk.Entry(frame_форма, width=25)
        self.цвет.grid(row=3, column=1, padx=5, pady=2)

        # Материал
        ttk.Label(frame_форма, text="Материал:").grid(row=4, column=0, padx=5, pady=2, sticky='w')
        self.материал = ttk.Entry(frame_форма, width=25)
        self.материал.grid(row=4, column=1, padx=5, pady=2)

        # Цена
        ttk.Label(frame_форма, text="Цена:").grid(row=5, column=0, padx=5, pady=2, sticky='w')
        self.цена = ttk.Entry(frame_форма, width=25)
        self.цена.grid(row=5, column=1, padx=5, pady=2)

        # Изображение
        ttk.Label(frame_форма, text="Изображение:").grid(row=6, column=0, padx=5, pady=2, sticky='w')
        self.путь_к_изображению = ttk.Entry(frame_форма, state='readonly', width=25)
        self.путь_к_изображению.grid(row=6, column=1, padx=5, pady=2)
        ttk.Button(frame_форма, text="Выбрать", command=self.выбрать_изображение).grid(row=6, column=2, padx=5, pady=2)

        # Фурнитура
        frame_фурнитура = ttk.LabelFrame(frame_форма, text="Фурнитура", padding=5)
        frame_фурнитура.grid(row=7, column=0, columnspan=3, pady=5, sticky='ew')

        # Список фурнитуры
        self.список_фурнитуры_одежды = ttk.Treeview(frame_фурнитура, columns=("Тип", "Количество"), show="headings", height=2)
        self.список_фурнитуры_одежды.grid(row=0, column=0, columnspan=2, pady=2)
        self.список_фурнитуры_одежды.heading("Тип", text="Тип")
        self.список_фурнитуры_одежды.heading("Количество", text="Количество")
        self.список_фурнитуры_одежды.column("Тип", width=120)
        self.список_фурнитуры_одежды.column("Количество", width=80)

        # Поля для добавления фурнитуры
        ttk.Label(frame_фурнитура, text="Тип фурнитуры:").grid(row=1, column=0, padx=5, pady=2)
        self.тип_фурнитуры = ttk.Combobox(frame_фурнитура, values=[
            "Пуговицы", "Молнии", "Кнопки", "Крючки", "Петли",
            "Пряжки", "Застежки", "Липучки", "Шнурки", "Декоративные элементы"
        ], width=15)
        self.тип_фурнитуры.grid(row=1, column=1, padx=5, pady=2)

        ttk.Label(frame_фурнитура, text="Количество:").grid(row=2, column=0, padx=5, pady=2)
        self.количество_фурнитуры = ttk.Entry(frame_фурнитура, width=15)
        self.количество_фурнитуры.grid(row=2, column=1, padx=5, pady=2)

        # Кнопки управления фурнитурой
        frame_кнопки_фурнитуры = ttk.Frame(frame_фурнитура)
        frame_кнопки_фурнитуры.grid(row=3, column=0, columnspan=2, pady=2)
        ttk.Button(frame_кнопки_фурнитуры, text="Добавить", command=self.добавить_фурнитуру_к_одежде).pack(side='left', padx=2)
        ttk.Button(frame_кнопки_фурнитуры, text="Удалить", command=self.удалить_фурнитуру_из_одежды).pack(side='left', padx=2)

        # Трикотажные полотна
        frame_трикотаж = ttk.LabelFrame(frame_форма, text="Трикотажные полотна", padding=5)
        frame_трикотаж.grid(row=8, column=0, columnspan=3, pady=5, sticky='ew')

        # Список трикотажных полотен
        self.список_трикотажа_одежды = ttk.Treeview(frame_трикотаж, columns=("Тип", "Количество"), show="headings", height=2)
        self.список_трикотажа_одежды.grid(row=0, column=0, columnspan=2, pady=2)
        self.список_трикотажа_одежды.heading("Тип", text="Тип")
        self.список_трикотажа_одежды.heading("Количество", text="Количество")
        self.список_трикотажа_одежды.column("Тип", width=120)
        self.список_трикотажа_одежды.column("Количество", width=80)

        # Поля для добавления трикотажного полотна
        ttk.Label(frame_трикотаж, text="Тип полотна:").grid(row=1, column=0, padx=5, pady=2)
        self.тип_трикотажа = ttk.Combobox(frame_трикотаж, values=[
            "Кулир", "Футер", "Начес", "Вискоза", "Рибана",
            "Эластан", "Модал", "Интерлок", "Капитоний", "Бифлекс"
        ], width=15)
        self.тип_трикотажа.grid(row=1, column=1, padx=5, pady=2)

        ttk.Label(frame_трикотаж, text="Количество:").grid(row=2, column=0, padx=5, pady=2)
        self.количество_трикотажа = ttk.Entry(frame_трикотаж, width=15)
        self.количество_трикотажа.grid(row=2, column=1, padx=5, pady=2)

        # Кнопки управления трикотажными полотнами
        frame_кнопки_трикотажа = ttk.Frame(frame_трикотаж)
        frame_кнопки_трикотажа.grid(row=3, column=0, columnspan=2, pady=2)
        ttk.Button(frame_кнопки_трикотажа, text="Добавить", command=self.добавить_трикотаж_к_одежде).pack(side='left', padx=2)
        ttk.Button(frame_кнопки_трикотажа, text="Удалить", command=self.удалить_трикотаж_из_одежды).pack(side='left', padx=2)

        # Кнопка добавления
        ttk.Button(frame_форма, text="Добавить одежду", command=self.добавить_одежду).grid(row=9, column=0, columnspan=3, pady=5)

        # Правая часть - предпросмотр изображения
        frame_предпросмотр = ttk.LabelFrame(main_frame, text="Предпросмотр", padding=5)
        frame_предпросмотр.pack(side='right', fill='both', expand=True, padx=5, pady=5)

        self.метка_предпросмотра = ttk.Label(frame_предпросмотр)
        self.метка_предпросмотра.pack(expand=True)

    def добавить_фурнитуру_к_одежде(self):
        тип = self.тип_фурнитуры.get()
        количество = self.количество_фурнитуры.get()

        if not тип or not количество:
            messagebox.showwarning("Предупреждение", "Заполните все поля")
            return

        try:
            количество = int(количество)
            if количество <= 0:
                raise ValueError
        except ValueError:
            messagebox.showerror("Ошибка", "Количество должно быть положительным числом")
            return

        # Проверяем, нет ли уже такой фурнитуры в списке
        for item in self.список_фурнитуры_одежды.get_children():
            if self.список_фурнитуры_одежды.item(item)['values'][0] == тип:
                messagebox.showwarning("Предупреждение", "Этот тип фурнитуры уже добавлен")
                return

        # Добавляем фурнитуру в список
        self.список_фурнитуры_одежды.insert("", "end", values=(тип, количество))

        # Очищаем поля
        self.тип_фурнитуры.set('')
        self.количество_фурнитуры.delete(0, tk.END)

    def удалить_фурнитуру_из_одежды(self):
        selected_item = self.список_фурнитуры_одежды.selection()
        if not selected_item:
            messagebox.showwarning("Предупреждение", "Выберите фурнитуру для удаления")
            return

        self.список_фурнитуры_одежды.delete(selected_item)

    def добавить_трикотаж_к_одежде(self):
        тип = self.тип_трикотажа.get()
        количество = self.количество_трикотажа.get()

        if not тип or not количество:
            messagebox.showwarning("Предупреждение", "Заполните все поля")
            return

        try:
            количество = int(количество)
            if количество <= 0:
                raise ValueError
        except ValueError:
            messagebox.showerror("Ошибка", "Количество должно быть положительным числом")
            return

        # Проверяем, нет ли уже такого полотна в списке
        for item in self.список_трикотажа_одежды.get_children():
            if self.список_трикотажа_одежды.item(item)['values'][0] == тип:
                messagebox.showwarning("Предупреждение", "Этот тип полотна уже добавлен")
                return

        # Добавляем полотно в список
        self.список_трикотажа_одежды.insert("", "end", values=(тип, количество))

        # Очищаем поля
        self.тип_трикотажа.set('')
        self.количество_трикотажа.delete(0, tk.END)

    def удалить_трикотаж_из_одежды(self):
        selected_item = self.список_трикотажа_одежды.selection()
        if not selected_item:
            messagebox.showwarning("Предупреждение", "Выберите полотно для удаления")
            return

        self.список_трикотажа_одежды.delete(selected_item)

    def сгенерировать_штрих_код(self):
        """Генерирует уникальный 13-значный штрих-код"""
        while True:
            # Генерируем 12 случайных цифр
            barcode = ''.join(random.choices(string.digits, k=12))

            # Вычисляем контрольную сумму (алгоритм EAN-13)
            checksum = 0
            for i in range(12):
                if i % 2 == 0:
                    checksum += int(barcode[i])
                else:
                    checksum += int(barcode[i]) * 3

            # Вычисляем последнюю цифру
            checksum = (10 - (checksum % 10)) % 10
            barcode = barcode + str(checksum)

            # Проверяем, не существует ли уже такой штрих-код
            self.cursor.execute("SELECT COUNT(*) FROM clothing WHERE barcode = %s", (barcode,))
            if self.cursor.fetchone()[0] == 0:
                return barcode

    def добавить_одежду(self):
        try:
            тип = self.тип_одежды.get()
            название = self.название.get()
            размер = self.размер.get()
            цвет = self.цвет.get()
            материал = self.материал.get()
            цена = float(self.цена.get())
            путь_к_изображению = self.путь_к_изображению.get()
            штрих_код = self.сгенерировать_штрих_код()

            # Проверяем наличие достаточного количества фурнитуры
            for item in self.список_фурнитуры_одежды.get_children():
                тип_фурнитуры, количество = self.список_фурнитуры_одежды.item(item)['values']
                # Проверяем наличие фурнитуры в базе
                self.cursor.execute("""
                    SELECT SUM(quantity)
                    FROM received_fittings
                    WHERE type = %s AND status = 'Доставлено'
                """, (тип_фурнитуры,))
                доступное_количество = self.cursor.fetchone()[0] or 0
                if доступное_количество < количество:
                    raise ValueError(f"Недостаточно фурнитуры типа '{тип_фурнитуры}'. Доступно: {доступное_количество}, требуется: {количество}")

            # Проверяем наличие достаточного количества трикотажного полотна
            for item in self.список_трикотажа_одежды.get_children():
                тип_полотна, количество = self.список_трикотажа_одежды.item(item)['values']
                # Проверяем наличие полотна в базе
                self.cursor.execute("""
                    SELECT SUM(quantity)
                    FROM received_fittings
                    WHERE type = %s AND status = 'Доставлено'
                """, (тип_полотна,))
                доступное_количество = self.cursor.fetchone()[0] or 0
                if доступное_количество < количество:
                    raise ValueError(f"Недостаточно трикотажного полотна типа '{тип_полотна}'. Доступно: {доступное_количество}, требуется: {количество}")

            if тип == "Верхняя одежда":
                одежда = Outerwear(название, размер, цвет, материал, цена, путь_к_изображению)
            elif тип == "Нижнее белье":
                одежда = Underwear(название, размер, цвет, материал, цена, путь_к_изображению)
            elif тип == "Штаны/Шорты":
                одежда = PantsShorts(название, размер, цвет, материал, цена, путь_к_изображению)
            elif тип == "Верхняя нательная одежда":
                одежда = Top(название, размер, цвет, материал, цена, путь_к_изображению)
            elif тип == "Комплекты муж":
                одежда = MensSet(название, размер, цвет, материал, цена, путь_к_изображению)
            elif тип == "Комплекты жен":
                одежда = WomensSet(название, размер, цвет, материал, цена, путь_к_изображению)
            else:
                raise ValueError("Неверный тип одежды")

            # Добавляем одежду в базу данных
            self.cursor.execute("""
                INSERT INTO clothing (name, type, size, color, material, price, image_path, barcode)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """, (название, тип, размер, цвет, материал, цена, путь_к_изображению, штрих_код))

            clothing_id = self.cursor.fetchone()[0]

            # Добавляем фурнитуру для одежды и вычитаем из запасов
            for item in self.список_фурнитуры_одежды.get_children():
                тип_фурнитуры, количество = self.список_фурнитуры_одежды.item(item)['values']
                # Добавляем связь с одеждой
                self.cursor.execute("""
                    INSERT INTO clothing_fittings (clothing_id, fitting_type, quantity)
                    VALUES (%s, %s, %s)
                """, (clothing_id, тип_фурнитуры, количество))

                # Вычитаем из запасов фурнитуры (только из самой старой записи)
                self.cursor.execute("""
                    SELECT id, quantity
                    FROM received_fittings
                    WHERE type = %s AND status = 'Доставлено' AND quantity > 0
                    ORDER BY received_at ASC
                    LIMIT 1
                """, (тип_фурнитуры,))
                row = self.cursor.fetchone()
                if row:
                    fitting_id, fitting_qty = row
                    new_qty = fitting_qty - количество if fitting_qty > количество else 0
                    self.cursor.execute("""
                        UPDATE received_fittings
                        SET quantity = %s
                        WHERE id = %s
                    """, (new_qty, fitting_id))

            # Добавляем трикотажные полотна для одежды и вычитаем из запасов
            for item in self.список_трикотажа_одежды.get_children():
                тип_полотна, количество = self.список_трикотажа_одежды.item(item)['values']
                # Добавляем связь с одеждой
                self.cursor.execute("""
                    INSERT INTO clothing_knit_fabrics (clothing_id, fabric_type, quantity)
                    VALUES (%s, %s, %s)
                """, (clothing_id, тип_полотна, количество))

                # Вычитаем из запасов трикотажного полотна (только из самой старой записи)
                self.cursor.execute("""
                    SELECT id, quantity
                    FROM received_fittings
                    WHERE type = %s AND status = 'Доставлено' AND quantity > 0
                    ORDER BY received_at ASC
                    LIMIT 1
                """, (тип_полотна,))
                row = self.cursor.fetchone()
                if row:
                    fabric_id, fabric_qty = row
                    new_qty = fabric_qty - количество if fabric_qty > количество else 0
                    self.cursor.execute("""
                        UPDATE received_fittings
                        SET quantity = %s
                        WHERE id = %s
                    """, (new_qty, fabric_id))

            self.conn.commit()
            self.фабрика.add_clothing(одежда)
            self.обновить_список_одежды()
            messagebox.showinfo("Успех", "Одежда успешно добавлена")

            # Очистка полей
            self.название.delete(0, tk.END)
            self.размер.set('')
            self.цвет.delete(0, tk.END)
            self.материал.delete(0, tk.END)
            self.цена.delete(0, tk.END)
            self.путь_к_изображению.configure(state='normal')
            self.путь_к_изображению.delete(0, tk.END)
            self.путь_к_изображению.configure(state='readonly')
            self.метка_предпросмотра.configure(image='')

            # Очистка списков фурнитуры и трикотажа
            for item in self.список_фурнитуры_одежды.get_children():
                self.список_фурнитуры_одежды.delete(item)
            for item in self.список_трикотажа_одежды.get_children():
                self.список_трикотажа_одежды.delete(item)

        except ValueError as e:
            messagebox.showerror("Ошибка", str(e))
        except Error as e:
            self.conn.rollback()
            messagebox.showerror("Ошибка базы данных", str(e))

    def создать_интерфейс_склада(self):
        # Основной контейнер
        main_frame = ttk.Frame(self.frame_склад)
        main_frame.pack(expand=True, fill='both', padx=20, pady=20)

        # Фильтры
        frame_фильтры = ttk.LabelFrame(main_frame, text="Фильтры", padding=10)
        frame_фильтры.pack(fill='x', padx=5, pady=5)

        ttk.Label(frame_фильтры, text="Тип:").pack(side='left', padx=5)
        self.фильтр_тип = ttk.Combobox(frame_фильтры, values=[
            "Верхняя одежда",
            "Нижнее белье",
            "Штаны/Шорты",
            "Верхняя нательная одежда",
            "Комплекты муж",
            "Комплекты жен"
        ], width=20)
        self.фильтр_тип.pack(side='left', padx=5)

        ttk.Button(frame_фильтры, text="Применить фильтр", command=self.обновить_список_одежды).pack(side='left', padx=5)

        # Кнопка экспорта в Excel
        ttk.Button(main_frame, text="Экспорт в Excel", command=self.экспорт_склада_excel).pack(pady=5)

        # Список одежды и изображение
        frame_содержимое = ttk.Frame(main_frame)
        frame_содержимое.pack(expand=True, fill='both', padx=5, pady=5)

        # Список одежды
        frame_список = ttk.LabelFrame(frame_содержимое, text="Список одежды", padding=10)
        frame_список.pack(side='left', fill='both', expand=True, padx=5, pady=5)

        self.список_одежды = ttk.Treeview(frame_список, columns=("Название", "Размер", "Цвет", "Материал", "Цена"), show="headings", selectmode='extended')
        self.список_одежды.pack(expand=True, fill='both')

        for col in ("Название", "Размер", "Цвет", "Материал", "Цена"):
            self.список_одежды.heading(col, text=col)
            self.список_одежды.column(col, width=100)

        # Кнопка переноса в ассортимент
        ttk.Button(frame_список, text="Перенести в ассортимент магазина", command=self.перенести_в_ассортимент).pack(fill='x', pady=5)

        # Изображение
        frame_изображение = ttk.LabelFrame(frame_содержимое, text="Изображение", padding=10)
        frame_изображение.pack(side='right', fill='both', expand=True, padx=5, pady=5)

        self.метка_изображения = ttk.Label(frame_изображение)
        self.метка_изображения.pack(expand=True)

        # Общая стоимость
        self.метка_стоимости = ttk.Label(main_frame, text="", font=('Arial', 12, 'bold'))
        self.метка_стоимости.pack(pady=10)

        # Привязываем событие выбора элемента в списке
        self.список_одежды.bind('<<TreeviewSelect>>', self.показать_изображение)

        self.обновить_список_одежды()

    def перенести_в_ассортимент(self):
        selected_items = self.список_одежды.selection()
        if not selected_items:
            messagebox.showwarning("Предупреждение", "Выберите одежду для переноса")
            return
        count = 0
        for item_id in selected_items:
            item = self.список_одежды.item(item_id)
            values = item['values']
            # Найти id по названию (и другим параметрам для уникальности)
            self.cursor.execute("""
                SELECT id FROM clothing WHERE name = %s AND size = %s AND color = %s AND material = %s AND price = %s AND (material IS NULL OR material != 'Ассортимент магазина')
                LIMIT 1
            """, (values[0], values[1], values[2], values[3], float(values[4].replace(' руб.', ''))))
            row = self.cursor.fetchone()
            if not row:
                continue
            clothing_id = row[0]
            # Обновить поле material
            self.cursor.execute("""
                UPDATE clothing SET material = 'Ассортимент магазина' WHERE id = %s
            """, (clothing_id,))
            count += 1
        self.conn.commit()
        self.обновить_список_одежды()
        self.обновить_список_готовых_изделий()
        messagebox.showinfo("Успех", f"{count} вещ(ей) успешно перенесено в ассортимент магазина")

    def показать_изображение(self, event):
        selected_item = self.список_одежды.selection()
        if selected_item:
            item = self.список_одежды.item(selected_item[0])
            values = item['values']
            название = values[0]

            # Находим одежду по названию
            одежда = next((о for о in self.фабрика.inventory if о.name == название), None)
            if одежда and одежда.image_path:
                try:
                    image = Image.open(одежда.image_path)
                    image = image.resize((300, 300), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(image)
                    self.метка_изображения.configure(image=photo)
                    self.метка_изображения.image = photo
                except Exception as e:
                    self.метка_изображения.configure(image='')
                    messagebox.showerror("Ошибка", f"Не удалось загрузить изображение: {e}")
            else:
                self.метка_изображения.configure(image='')

    def обновить_список_одежды(self):
        # Очищаем список
        for item in self.список_одежды.get_children():
            self.список_одежды.delete(item)

        # Получаем значение фильтра
        выбранный_тип = self.фильтр_тип.get()

        # Формируем SQL-запрос
        query = """
            SELECT name, size, color, material, price
            FROM clothing
            WHERE (material IS NULL OR material != 'Ассортимент магазина')
        """
        params = []
        if выбранный_тип and выбранный_тип != "":
            query += " AND type = %s"
            params.append(выбранный_тип)
        query += " ORDER BY created_at DESC"

        self.cursor.execute(query, params)
        rows = self.cursor.fetchall()

        for row in rows:
            self.список_одежды.insert("", "end", values=(
                row[0],  # Название
                row[1],  # Размер
                row[2],  # Цвет
                row[3],  # Материал
                f"{row[4]} руб."
            ))

        # Обновляем общую стоимость
        общая_стоимость = sum(row[4] for row in rows if row[4] is not None)
        self.метка_стоимости.configure(text=f"Общая стоимость: {общая_стоимость} руб.")

    def выбрать_изображение(self):
        путь = filedialog.askopenfilename(
            title="Выберите изображение",
            filetypes=[("Изображения", "*.png *.jpg *.jpeg *.bmp *.gif")]
        )
        if путь:
            self.путь_к_изображению.configure(state='normal')
            self.путь_к_изображению.delete(0, tk.END)
            self.путь_к_изображению.insert(0, путь)
            self.путь_к_изображению.configure(state='readonly')

            # Показываем предпросмотр
            try:
                image = Image.open(путь)
                image = image.resize((150, 150), Image.Resampling.LANCZOS)  # Уменьшаем размер предпросмотра
                photo = ImageTk.PhotoImage(image)
                self.метка_предпросмотра.configure(image=photo)
                self.метка_предпросмотра.image = photo
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось загрузить изображение: {e}")

    def создать_интерфейс_фурнитуры(self):
        # Основной контейнер
        main_frame = ttk.Frame(self.frame_фурнитура)
        main_frame.pack(expand=True, fill='both', padx=20, pady=20)

        # Фильтры
        frame_фильтры = ttk.LabelFrame(main_frame, text="Фильтры", padding=10)
        frame_фильтры.pack(fill='x', padx=5, pady=5)

        ttk.Label(frame_фильтры, text="Статус:").pack(side='left', padx=5)
        self.фильтр_статус = ttk.Combobox(frame_фильтры, values=[
            "Все",
            "В обработке",
            "В пути",
            "Доставлено",
            "Отменено"
        ], width=20)
        self.фильтр_статус.pack(side='left', padx=5)
        self.фильтр_статус.set("Все")

        ttk.Button(frame_фильтры, text="Применить фильтр", command=self.обновить_список_фурнитуры).pack(side='left', padx=5)

        # Кнопка управления статусами
        ttk.Button(frame_фильтры, text="Управление статусами", command=self.управление_статусами).pack(side='left', padx=5)

        # Список фурнитуры
        frame_список = ttk.LabelFrame(main_frame, text="Список фурнитуры", padding=10)
        frame_список.pack(expand=True, fill='both', padx=5, pady=5)

        self.список_фурнитуры = ttk.Treeview(frame_список, columns=(
            "Название", "Тип", "Количество", "Цена", "Статус", "Дата заказа", "Поставщик"
        ), show="headings")
        self.список_фурнитуры.pack(expand=True, fill='both')

        for col in ("Название", "Тип", "Количество", "Цена", "Статус", "Дата заказа", "Поставщик"):
            self.список_фурнитуры.heading(col, text=col)
            self.список_фурнитуры.column(col, width=100)

        # Обновляем список фурнитуры при создании интерфейса
        self.обновить_список_фурнитуры()

    def управление_статусами(self):
        # Создаем окно управления статусами
        окно_статусов = tk.Toplevel(self.frame_фурнитура)
        окно_статусов.title("Управление статусами заказов")
        окно_статусов.geometry("800x600")

        # Список заказов
        frame_список = ttk.LabelFrame(окно_статусов, text="Заказы", padding=10)
        frame_список.pack(fill='both', expand=True, padx=10, pady=5)

        # Создаем Treeview для отображения заказов
        columns = ("ID", "Тип", "Количество", "Цена", "Статус", "Дата заказа", "Поставщик")
        список_заказов = ttk.Treeview(frame_список, columns=columns, show="headings")
        список_заказов.pack(fill='both', expand=True)

        # Настраиваем заголовки
        for col in columns:
            список_заказов.heading(col, text=col)
            список_заказов.column(col, width=100)

        # Загружаем данные о заказах
        self.cursor.execute("""
            SELECT
                fo.id,
                fo.type,
                fo.quantity,
                fo.price,
                fo.status,
                fo.created_at,
                s.company_name
            FROM fittings_orders fo
            JOIN suppliers s ON fo.supplier_id = s.id
            ORDER BY fo.created_at DESC
        """)

        for заказ in self.cursor.fetchall():
            список_заказов.insert("", "end", values=(
                заказ[0],  # ID
                заказ[1],  # Тип
                заказ[2],  # Количество
                f"{заказ[3]} руб." if заказ[3] else "Не указана",  # Цена
                заказ[4],  # Статус
                заказ[5].strftime("%Y-%m-%d %H:%M"),  # Дата заказа
                заказ[6]   # Поставщик
            ))

        # Фрейм для изменения статуса
        frame_статус = ttk.LabelFrame(окно_статусов, text="Изменение статуса", padding=10)
        frame_статус.pack(fill='x', padx=10, pady=5)

        ttk.Label(frame_статус, text="Новый статус:").pack(side='left', padx=5)
        новый_статус = ttk.Combobox(frame_статус, values=[
            "В обработке",
            "В пути",
            "Доставлено",
            "Отменено"
        ], width=20)
        новый_статус.pack(side='left', padx=5)

        def обновить_статус():
            selected_items = список_заказов.selection()
            if not selected_items:
                messagebox.showwarning("Предупреждение", "Выберите заказ для изменения статуса")
                return

            if not новый_статус.get():
                messagebox.showwarning("Предупреждение", "Выберите новый статус")
                return

            try:
                # Получаем ID выбранного заказа
                order_id = список_заказов.item(selected_items[0])['values'][0]

                # Обновляем статус заказа
                self.обновить_статус_заказа(order_id, новый_статус.get())

                # Обновляем список заказов
                for item in список_заказов.get_children():
                    if список_заказов.item(item)['values'][0] == order_id:
                        values = список_заказов.item(item)['values']
                        values = list(values)
                        values[4] = новый_статус.get()  # Обновляем статус
                        список_заказов.item(item, values=values)
                        break

                messagebox.showinfo("Успех", "Статус успешно обновлен")

            except Error as e:
                self.conn.rollback()
                messagebox.showerror("Ошибка базы данных", str(e))

        ttk.Button(frame_статус, text="Обновить статус", command=обновить_статус).pack(side='left', padx=5)

        # Делаем окно модальным
        окно_статусов.transient(self.frame_фурнитура)
        окно_статусов.grab_set()
        self.frame_фурнитура.wait_window(окно_статусов)

    def обновить_список_фурнитуры(self):
        # Очищаем список
        for item in self.список_фурнитуры.get_children():
            self.список_фурнитуры.delete(item)

        # Получаем выбранный статус
        выбранный_статус = self.фильтр_статус.get()

        # Базовый запрос
        запрос = """
            SELECT
                fo.type as name,
                fo.type,
                fo.quantity,
                fo.price,
                fo.status,
                fo.created_at,
                s.company_name as supplier
            FROM fittings_orders fo
            JOIN suppliers s ON fo.supplier_id = s.id
            WHERE 1=1
        """

        # Добавляем фильтр по статусу, если выбран
        if выбранный_статус and выбранный_статус != "Все":
            запрос += " AND fo.status = %s"

        запрос += " ORDER BY fo.created_at DESC"

        # Выполняем запрос
        if выбранный_статус and выбранный_статус != "Все":
            self.cursor.execute(запрос, (выбранный_статус,))
        else:
            self.cursor.execute(запрос)

        for фурнитура in self.cursor.fetchall():
            self.список_фурнитуры.insert("", "end", values=(
                фурнитура[0],  # Название
                фурнитура[1],  # Тип
                фурнитура[2],  # Количество
                f"{фурнитура[3]} руб.",  # Цена
                фурнитура[4],  # Статус
                фурнитура[5].strftime("%Y-%m-%d %H:%M"),  # Дата заказа
                фурнитура[6]   # Поставщик
            ))

    def создать_интерфейс_готовых_изделий(self):
        # Основной контейнер
        main_frame = ttk.Frame(self.frame_готовые_изделия)
        main_frame.pack(expand=True, fill='both', padx=20, pady=20)

        # Кнопка открытия ассортимента магазина
        ttk.Button(main_frame, text="Открыть ассортимент магазина",
                  command=self.открыть_ассортимент_магазина).pack(pady=20)

    def открыть_ассортимент_магазина(self):
        АссортиментМагазинаGUI(self.root, self.cursor, self.conn)

    def экспорт_полной_документации(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if not file_path:
            return

        doc = Document()
        doc.add_heading("Документация по фабрике одежды", 0)

        # 1. Ассортимент магазина
        doc.add_heading("Ассортимент магазина", level=1)
        table_ассортимент = doc.add_table(rows=1, cols=5)
        hdr_cells = table_ассортимент.rows[0].cells
        hdr_cells[0].text = 'Тип'
        hdr_cells[1].text = 'Размер'
        hdr_cells[2].text = 'Цвет'
        hdr_cells[3].text = 'Количество'
        hdr_cells[4].text = 'Цена'

        self.cursor.execute("""
            SELECT type, size, color, quantity, price
            FROM clothing
            WHERE material = 'Ассортимент магазина'
            ORDER BY type, size
        """)
        total_ассортимент = 0
        for row in self.cursor.fetchall():
            row_cells = table_ассортимент.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
            total_ассортимент += row[3] * row[4] if row[3] and row[4] else 0

        doc.add_paragraph(f"Общая стоимость ассортимента: {total_ассортимент:.2f} руб.")

        # 2. Продажи
        doc.add_heading("История продаж", level=1)
        table_продажи = doc.add_table(rows=1, cols=7)
        hdr_cells = table_продажи.rows[0].cells
        hdr_cells[0].text = 'Тип'
        hdr_cells[1].text = 'Размер'
        hdr_cells[2].text = 'Цвет'
        hdr_cells[3].text = 'Количество'
        hdr_cells[4].text = 'Цена за ед.'
        hdr_cells[5].text = 'Общая сумма'
        hdr_cells[6].text = 'Дата продажи'

        self.cursor.execute("""
            SELECT type, size, color, quantity, price, total_price, sale_date
            FROM sales
            ORDER BY sale_date DESC
        """)
        total_продажи = 0
        for row in self.cursor.fetchall():
            row_cells = table_продажи.add_row().cells
            for i, value in enumerate(row):
                if i == 6:  # Дата продажи
                    row_cells[i].text = value.strftime("%Y-%m-%d %H:%M")
                else:
                    row_cells[i].text = str(value)
            total_продажи += row[5] if row[5] else 0

        doc.add_paragraph(f"Общая сумма продаж: {total_продажи:.2f} руб.")

        # 3. Склад фурнитуры
        doc.add_heading("Склад фурнитуры", level=1)
        table_фурнитура = doc.add_table(rows=1, cols=5)
        hdr_cells = table_фурнитура.rows[0].cells
        hdr_cells[0].text = 'Тип'
        hdr_cells[1].text = 'Категория'
        hdr_cells[2].text = 'Количество'
        hdr_cells[3].text = 'Цена'
        hdr_cells[4].text = 'Дата поступления'

        self.cursor.execute("""
            SELECT
                rf.type,
                fo.category,
                rf.quantity,
                rf.price,
                rf.received_at
            FROM received_fittings rf
            JOIN fittings_orders fo ON rf.order_id = fo.id
            WHERE rf.status = 'Доставлено'
            ORDER BY rf.received_at DESC
        """)
        total_фурнитура = 0
        for row in self.cursor.fetchall():
            row_cells = table_фурнитура.add_row().cells
            for i, value in enumerate(row):
                if i == 4:  # Дата поступления
                    row_cells[i].text = value.strftime("%Y-%m-%d %H:%M")
                else:
                    row_cells[i].text = str(value)
            total_фурнитура += row[2] * row[3] if row[2] and row[3] else 0

        doc.add_paragraph(f"Общая стоимость фурнитуры на складе: {total_фурнитура:.2f} руб.")

        # Добавляем итоговую информацию
        doc.add_heading("Итоговая информация", level=1)
        doc.add_paragraph(f"Общая стоимость ассортимента: {total_ассортимент:.2f} руб.")
        doc.add_paragraph(f"Общая сумма продаж: {total_продажи:.2f} руб.")
        doc.add_paragraph(f"Общая стоимость фурнитуры на складе: {total_фурнитура:.2f} руб.")
        doc.add_paragraph(f"Общая стоимость активов: {total_ассортимент + total_фурнитура:.2f} руб.")

        # Сохраняем документ
        doc.save(file_path)
        messagebox.showinfo("Экспорт завершен", f"Документация успешно экспортирована в {file_path}")

    def обновить_список_готовых_изделий(self):
        # Очищаем список
        for item in self.список_готовых.get_children():
            self.список_готовых.delete(item)

        # Получаем значения фильтров
        выбранный_тип = self.фильтр_тип_ассортимент.get()
        выбранный_размер = self.фильтр_размер_ассортимент.get()
        выбранный_цвет = self.фильтр_цвет_ассортимент.get()

        # Формируем SQL-запрос
        query = """
            SELECT id, barcode, name, type, size, color, quantity, price
            FROM clothing
            WHERE material = 'Ассортимент магазина'
        """
        params = []

        # Добавляем условия фильтрации
        if выбранный_тип and выбранный_тип != "Все":
            query += " AND type = %s"
            params.append(выбранный_тип)
        if выбранный_размер and выбранный_размер != "Все":
            query += " AND size = %s"
            params.append(выбранный_размер)
        if выбранный_цвет and выбранный_цвет != "Все":
            query += " AND color = %s"
            params.append(выбранный_цвет)

        query += " ORDER BY created_at DESC"

        # Выполняем запрос
        self.cursor.execute(query, params)

        total = 0
        for изделие in self.cursor.fetchall():
            self.список_готовых.insert("", "end", values=(
                изделие[0],  # ID
                изделие[1],  # Штрих-код
                изделие[2],  # Название
                изделие[3],  # Тип
                изделие[4],  # Размер
                изделие[5],  # Цвет
                изделие[6],  # Количество
                f"{изделие[7]} руб."  # Цена
            ))
            try:
                qty = int(изделие[6]) if изделие[6] is not None else 1
                price = float(изделие[7]) if изделие[7] is not None else 0
                total += qty * price
            except Exception:
                pass

        self.метка_стоимости_ассортимента.configure(text=f"Общая стоимость ассортимента: {total:.2f} руб.")

    def добавить_тестовые_данные(self):
        try:
            # Добавляем тестовую одежду
            одежда = [
                ("Пальто зимнее", "Верхняя одежда", "L", "Черный", "Шерсть", 15000, "images/palto.jpg"),
                ("Футболка базовая", "Повседневная", "M", "Белый", "Хлопок", 2000, "images/futbolka.jpg"),
                ("Джинсы классические", "Повседневная", "42", "Синий", "Джинс", 5000, "images/djinsi.jpg")
            ]

            for item in одежда:
                self.cursor.execute("""
                    INSERT INTO clothing (name, type, size, color, material, price, image_path)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, item)

            # Добавляем тестовую фурнитуру и трикотаж
            фурнитура_и_трикотаж = [
                # Фурнитура
                ("Пуговицы", 1000, 5.0, "Доставлено"),
                ("Молнии", 500, 15.0, "Доставлено"),
                ("Кнопки", 2000, 3.0, "Доставлено"),
                ("Крючки", 800, 4.0, "Доставлено"),
                ("Петли", 600, 6.0, "Доставлено"),
                # Трикотажные полотна
                ("Кулир", 100, 200.0, "Доставлено"),
                ("Футер", 80, 250.0, "Доставлено"),
                ("Начес", 60, 300.0, "Доставлено"),
                ("Вискоза", 90, 280.0, "Доставлено"),
                ("Рибана", 70, 220.0, "Доставлено")
            ]

            for item in фурнитура_и_трикотаж:
                self.cursor.execute("""
                    INSERT INTO received_fittings (type, quantity, price, status)
                    VALUES (%s, %s, %s, %s)
                """, item)

            self.conn.commit()
            messagebox.showinfo("Успех", "Тестовые данные успешно добавлены в базу данных")

            # Обновляем списки в интерфейсе
            self.обновить_список_одежды()
            self.обновить_список_склада_фурнитуры()

        except Error as e:
            self.conn.rollback()
            messagebox.showerror("Ошибка", f"Ошибка при добавлении тестовых данных: {str(e)}")

    def создать_интерфейс_поставщиков(self):
        # Основной контейнер
        main_frame = ttk.Frame(self.frame_поставщики)
        main_frame.pack(expand=True, fill='both', padx=20, pady=20)

        # Левая часть - форма добавления поставщика
        frame_форма = ttk.LabelFrame(main_frame, text="Добавить поставщика", padding=10)
        frame_форма.pack(side='left', fill='both', expand=True, padx=5, pady=5)

        # Название компании
        ttk.Label(frame_форма, text="Название компании:").grid(row=0, column=0, padx=5, pady=5, sticky='w')
        self.название_компании = ttk.Entry(frame_форма, width=30)
        self.название_компании.grid(row=0, column=1, padx=5, pady=5)

        # Контактное лицо
        ttk.Label(frame_форма, text="Контактное лицо:").grid(row=1, column=0, padx=5, pady=5, sticky='w')
        self.контактное_лицо = ttk.Entry(frame_форма, width=30)
        self.контактное_лицо.grid(row=1, column=1, padx=5, pady=5)

        # Телефон
        ttk.Label(frame_форма, text="Телефон:").grid(row=2, column=0, padx=5, pady=5, sticky='w')
        self.телефон = ttk.Entry(frame_форма, width=30)
        self.телефон.grid(row=2, column=1, padx=5, pady=5)

        # Email
        ttk.Label(frame_форма, text="Email:").grid(row=3, column=0, padx=5, pady=5, sticky='w')
        self.email = ttk.Entry(frame_форма, width=30)
        self.email.grid(row=3, column=1, padx=5, pady=5)

        # Адрес
        ttk.Label(frame_форма, text="Адрес:").grid(row=4, column=0, padx=5, pady=5, sticky='w')
        self.адрес = ttk.Entry(frame_форма, width=30)
        self.адрес.grid(row=4, column=1, padx=5, pady=5)

        # Кнопка добавления
        ttk.Button(frame_форма, text="Добавить", command=self.добавить_поставщика).grid(row=5, column=0, columnspan=2, pady=10)

        # Правая часть - список поставщиков
        frame_список = ttk.LabelFrame(main_frame, text="Список поставщиков", padding=10)
        frame_список.pack(side='right', fill='both', expand=True, padx=5, pady=5)

        # Кнопки управления
        frame_кнопки = ttk.Frame(frame_список)
        frame_кнопки.pack(fill='x', pady=5)
        ttk.Button(frame_кнопки, text="Сделать заказ", command=self.сделать_заказ_фурнитуры).pack(side='left', padx=5)

        self.список_поставщиков = ttk.Treeview(frame_список, columns=("Компания", "Контакт", "Телефон", "Email", "Адрес"), show="headings")
        self.список_поставщиков.pack(expand=True, fill='both')

        for col in ("Компания", "Контакт", "Телефон", "Email", "Адрес"):
            self.список_поставщиков.heading(col, text=col)
            self.список_поставщиков.column(col, width=100)

        # Обновляем список поставщиков
        self.обновить_список_поставщиков()

    def добавить_поставщика(self):
        try:
            название = self.название_компании.get()
            контакт = self.контактное_лицо.get()
            телефон = self.телефон.get()
            email = self.email.get()
            адрес = self.адрес.get()

            if not все_заполнено([название, контакт, телефон, email, адрес]):
                raise ValueError("Все поля должны быть заполнены")

            # Добавляем поставщика в базу данных
            self.cursor.execute("""
                INSERT INTO suppliers (company_name, contact_person, phone, email, address)
                VALUES (%s, %s, %s, %s, %s)
                RETURNING id
            """, (название, контакт, телефон, email, адрес))

            supplier_id = self.cursor.fetchone()[0]
            self.conn.commit()

            # Добавляем поставщика в список
            self.список_поставщиков.insert("", "end", values=(
                название,
                контакт,
                телефон,
                email,
                адрес
            ), tags=(str(supplier_id),))  # ID поставщика в тегах

            # Очищаем поля
            self.название_компании.delete(0, tk.END)
            self.контактное_лицо.delete(0, tk.END)
            self.телефон.delete(0, tk.END)
            self.email.delete(0, tk.END)
            self.адрес.delete(0, tk.END)

            messagebox.showinfo("Успех", "Поставщик успешно добавлен")

        except ValueError as e:
            messagebox.showerror("Ошибка", str(e))
        except Error as e:
            self.conn.rollback()
            messagebox.showerror("Ошибка базы данных", str(e))

    def обновить_список_поставщиков(self):
        # Очищаем список
        for item in self.список_поставщиков.get_children():
            self.список_поставщиков.delete(item)

        # Получаем поставщиков из базы данных
        self.cursor.execute("""
            SELECT id, company_name, contact_person, phone, email, address
            FROM suppliers
            ORDER BY company_name
        """)

        for поставщик in self.cursor.fetchall():
            self.список_поставщиков.insert("", "end", values=(
                поставщик[1],  # Компания
                поставщик[2],  # Контакт
                поставщик[3],  # Телефон
                поставщик[4],  # Email
                поставщик[5]   # Адрес
            ), tags=(str(поставщик[0]),))  # ID поставщика в тегах

    def сделать_заказ_фурнитуры(self):
        # Получаем выбранного поставщика
        selected_item = self.список_поставщиков.selection()
        if not selected_item:
            messagebox.showwarning("Предупреждение", "Выберите поставщика")
            return

        # Получаем ID поставщика из тегов элемента
        supplier_id = int(self.список_поставщиков.item(selected_item[0])['tags'][0])
        поставщик = self.список_поставщиков.item(selected_item[0])['values']

        # Создаем окно заказа
        окно_заказа = tk.Toplevel(self.frame_поставщики)
        окно_заказа.title("Заказ фурнитуры")
        окно_заказа.geometry("500x400")

        # Отображаем информацию о поставщике
        ttk.Label(окно_заказа, text=f"Поставщик: {поставщик[0]}", font=('Arial', 10, 'bold')).pack(pady=5)
        ttk.Label(окно_заказа, text=f"Контакт: {поставщик[1]}").pack(pady=2)
        ttk.Label(окно_заказа, text=f"Телефон: {поставщик[2]}").pack(pady=2)

        # Форма заказа
        frame_заказ = ttk.LabelFrame(окно_заказа, text="Детали заказа", padding=10)
        frame_заказ.pack(fill='both', expand=True, padx=10, pady=10)

        # Категория товара
        ttk.Label(frame_заказ, text="Категория:").grid(row=0, column=0, padx=5, pady=5, sticky='w')
        категория = ttk.Combobox(frame_заказ, values=["Фурнитура", "Трикотажные полотна"], width=30)
        категория.grid(row=0, column=1, padx=5, pady=5)

        # Тип фурнитуры (будет меняться в зависимости от категории)
        ttk.Label(frame_заказ, text="Тип:").grid(row=1, column=0, padx=5, pady=5, sticky='w')
        тип = ttk.Combobox(frame_заказ, width=30)
        тип.grid(row=1, column=1, padx=5, pady=5)

        # Функция обновления списка типов в зависимости от категории
        def обновить_типы(event):
            if категория.get() == "Фурнитура":
                тип['values'] = [
                    "Пуговицы",
                    "Молнии",
                    "Кнопки",
                    "Крючки",
                    "Петли",
                    "Пряжки",
                    "Застежки",
                    "Липучки",
                    "Шнурки",
                    "Декоративные элементы"
                ]
            elif категория.get() == "Трикотажные полотна":
                тип['values'] = [
                    "Кулир",
                    "Футер",
                    "Начес",
                    "Вискоза",
                    "Рибана",
                    "Эластан",
                    "Модал",
                    "Интерлок",
                    "Капитоний",
                    "Бифлекс"
                ]
            тип.set('')  # Очищаем выбранное значение

        # Привязываем обновление типов к изменению категории
        категория.bind('<<ComboboxSelected>>', обновить_типы)

        # Цвет
        ttk.Label(frame_заказ, text="Цвет:").grid(row=2, column=0, padx=5, pady=5, sticky='w')
        цвет = ttk.Entry(frame_заказ, width=30)
        цвет.grid(row=2, column=1, padx=5, pady=5)

        # Количество
        ttk.Label(frame_заказ, text="Количество:").grid(row=3, column=0, padx=5, pady=5, sticky='w')
        количество = ttk.Entry(frame_заказ, width=30)
        количество.grid(row=3, column=1, padx=5, pady=5)

        # Цена за единицу
        ttk.Label(frame_заказ, text="Цена за единицу:").grid(row=4, column=0, padx=5, pady=5, sticky='w')
        цена = ttk.Entry(frame_заказ, width=30)
        цена.grid(row=4, column=1, padx=5, pady=5)

        # Дата доставки
        ttk.Label(frame_заказ, text="Дата доставки:").grid(row=5, column=0, padx=5, pady=5, sticky='w')
        дата = ttk.Entry(frame_заказ, width=30)
        дата.grid(row=5, column=1, padx=5, pady=5)

        # Комментарий
        ttk.Label(frame_заказ, text="Комментарий:").grid(row=6, column=0, padx=5, pady=5, sticky='w')
        комментарий = ttk.Entry(frame_заказ, width=30)
        комментарий.grid(row=6, column=1, padx=5, pady=5)

        def подтвердить_заказ():
            try:
                # Проверяем только категорию и тип
                if not категория.get() or not тип.get():
                    raise ValueError("Выберите категорию и тип")

                # Проверяем числовые значения только если они заполнены
                кол = int(количество.get()) if количество.get() else 0
                ц = float(цена.get()) if цена.get() else 0.0

                # Добавляем заказ в базу данных
                self.cursor.execute("""
                    INSERT INTO fittings_orders
                    (supplier_id, category, type, color, quantity, price, delivery_date, comment, status)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, 'В обработке')
                    RETURNING id
                """, (
                    supplier_id,
                    категория.get(),
                    тип.get(),
                    цвет.get() if цвет.get() else None,
                    кол if кол > 0 else None,
                    ц if ц > 0 else None,
                    дата.get() if дата.get() else None,
                    комментарий.get() if комментарий.get() else None
                ))

                order_id = self.cursor.fetchone()[0]

                # Создаем запись в received_fittings с начальным статусом "В обработке"
                self.cursor.execute("""
                    INSERT INTO received_fittings
                    (order_id, type, quantity, price, status)
                    VALUES (%s, %s, %s, %s, 'В обработке')
                """, (
                    order_id,
                    тип.get(),
                    кол if кол > 0 else 0,
                    ц if ц > 0 else 0.0
                ))

                self.conn.commit()

                # Создаем сообщение о заказе
                сообщение = f"""
                Заказ успешно оформлен!

                Поставщик: {поставщик[0]}
                Категория: {категория.get()}
                Тип: {тип.get()}
                Цвет: {цвет.get() if цвет.get() else 'Не указан'}
                Количество: {кол if кол > 0 else 'Не указано'}
                Цена за единицу: {ц if ц > 0 else 'Не указана'} руб.
                Общая стоимость: {кол * ц if кол > 0 and ц > 0 else 'Не указана'} руб.
                Дата доставки: {дата.get() if дата.get() else 'Не указана'}
                Комментарий: {комментарий.get() if комментарий.get() else 'Нет'}
                """

                messagebox.showinfo("Заказ оформлен", сообщение)
                окно_заказа.destroy()

                # Обновляем список фурнитуры
                self.обновить_список_фурнитуры()
                self.обновить_список_склада_фурнитуры()

            except ValueError as e:
                messagebox.showerror("Ошибка", str(e))
            except Error as e:
                self.conn.rollback()
                messagebox.showerror("Ошибка базы данных", str(e))

        # Кнопки
        frame_кнопки = ttk.Frame(frame_заказ)
        frame_кнопки.grid(row=7, column=0, columnspan=2, pady=10)

        ttk.Button(frame_кнопки, text="Подтвердить", command=подтвердить_заказ).pack(side='left', padx=5)
        ttk.Button(frame_кнопки, text="Отмена", command=окно_заказа.destroy).pack(side='left', padx=5)

    def создать_интерфейс_склада_фурнитуры(self):
        # Основной контейнер
        main_frame = ttk.Frame(self.frame_склад_фурнитуры)
        main_frame.pack(expand=True, fill='both', padx=20, pady=20)

        # Фильтры
        frame_фильтры = ttk.LabelFrame(main_frame, text="Фильтры", padding=10)
        frame_фильтры.pack(fill='x', padx=5, pady=5)

        ttk.Label(frame_фильтры, text="Категория:").pack(side='left', padx=5)
        self.фильтр_категория = ttk.Combobox(frame_фильтры, values=[
            "Все",
            "Фурнитура",
            "Трикотажные полотна"
        ], width=20)
        self.фильтр_категория.pack(side='left', padx=5)
        self.фильтр_категория.set("Все")

        ttk.Button(frame_фильтры, text="Применить фильтр",
                  command=self.обновить_список_склада_фурнитуры).pack(side='left', padx=5)

        # Список фурнитуры и трикотажа
        frame_список = ttk.LabelFrame(main_frame, text="Склад фурнитуры и трикотажа", padding=10)
        frame_список.pack(expand=True, fill='both', padx=5, pady=5)

        self.список_склада_фурнитуры = ttk.Treeview(frame_список, columns=(
            "Тип", "Категория", "Количество", "Цена", "Дата поступления"
        ), show="headings")
        self.список_склада_фурнитуры.pack(expand=True, fill='both')

        for col in ("Тип", "Категория", "Количество", "Цена", "Дата поступления"):
            self.список_склада_фурнитуры.heading(col, text=col)
            self.список_склада_фурнитуры.column(col, width=150)

        # Обновляем список при создании интерфейса
        self.обновить_список_склада_фурнитуры()

    def обновить_список_склада_фурнитуры(self):
        # Очищаем список
        for item in self.список_склада_фурнитуры.get_children():
            self.список_склада_фурнитуры.delete(item)

        # Получаем выбранную категорию
        выбранная_категория = self.фильтр_категория.get()

        # Базовый запрос
        запрос = """
            SELECT
                rf.type,
                fo.category,
                rf.quantity,
                rf.price,
                rf.received_at
            FROM received_fittings rf
            JOIN fittings_orders fo ON rf.order_id = fo.id
            WHERE rf.status = 'Доставлено'
        """

        # Добавляем фильтр по категории, если выбран
        if выбранная_категория and выбранная_категория != "Все":
            запрос += " AND fo.category = %s"

        запрос += " ORDER BY rf.received_at DESC"

        # Выполняем запрос
        if выбранная_категория and выбранная_категория != "Все":
            self.cursor.execute(запрос, (выбранная_категория,))
        else:
            self.cursor.execute(запрос)

        total_value = 0
        for item in self.cursor.fetchall():
            self.список_склада_фурнитуры.insert("", "end", values=(
                item[0],  # Тип
                item[1],  # Категория
                item[2],  # Количество
                f"{item[3]} руб.",  # Цена
                item[4].strftime("%Y-%m-%d %H:%M")  # Дата поступления
            ))
            # Добавляем к общей стоимости
            total_value += item[2] * item[3]  # количество * цена

        # Обновляем метку с общей стоимостью
        if hasattr(self, 'метка_стоимости_фурнитуры'):
            self.метка_стоимости_фурнитуры.configure(text=f"Общая стоимость фурнитуры на складе: {total_value:.2f} руб.")
        else:
            # Создаем метку, если она еще не существует
            self.метка_стоимости_фурнитуры = ttk.Label(self.frame_склад_фурнитуры, text=f"Общая стоимость фурнитуры на складе: {total_value:.2f} руб.", font=('Arial', 12, 'bold'))
            self.метка_стоимости_фурнитуры.pack(pady=10)

    def обновить_статус_заказа(self, order_id, новый_статус):
        try:
            # Обновляем статус в fittings_orders
            self.cursor.execute("""
                UPDATE fittings_orders
                SET status = %s
                WHERE id = %s
            """, (новый_статус, order_id))

            # Обновляем статус в received_fittings
            self.cursor.execute("""
                UPDATE received_fittings
                SET status = %s
                WHERE order_id = %s
            """, (новый_статус, order_id))

            self.conn.commit()

            # Обновляем списки
            self.обновить_список_фурнитуры()
            self.обновить_список_склада_фурнитуры()

        except Error as e:
            self.conn.rollback()
            messagebox.showerror("Ошибка базы данных", str(e))

    def экспорт_склада_excel(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if not file_path:
            return

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Склад одежды"

        headers = ["Название", "Размер", "Цвет", "Материал", "Цена"]
        ws.append(headers)

        # Получаем данные из базы
        self.cursor.execute("""
            SELECT name, size, color, material, price
            FROM clothing
            WHERE (material IS NULL OR material != 'Ассортимент магазина')
            ORDER BY created_at DESC
        """)
        for row in self.cursor.fetchall():
            ws.append(row)

        # Автоширина
        for col in ws.columns:
            max_length = 0
            col_letter = get_column_letter(col[0].column)
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws.column_dimensions[col_letter].width = max_length + 2

        wb.save(file_path)
        messagebox.showinfo("Экспорт завершен", f"Данные успешно экспортированы в {file_path}")

    def экспорт_ассортимента_word(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if not file_path:
            return

        doc = Document()
        doc.add_heading("Ассортимент магазина", 0)

        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Тип'
        hdr_cells[1].text = 'Размер'
        hdr_cells[2].text = 'Цвет'
        hdr_cells[3].text = 'Количество'
        hdr_cells[4].text = 'Цена'

        self.cursor.execute("""
            SELECT type, size, color, quantity, price
            FROM clothing
            WHERE material = 'Ассортимент магазина'
            ORDER BY created_at DESC
        """)
        for row in self.cursor.fetchall():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        doc.save(file_path)
        messagebox.showinfo("Экспорт завершен", f"Данные успешно экспортированы в {file_path}")

    def продать_одежду(self):
        selected_items = self.список_готовых.selection()
        if not selected_items:
            messagebox.showwarning("Предупреждение", "Выберите товар для продажи")
            return

        # Создаем окно продажи
        окно_продажи = tk.Toplevel(self.frame_готовые_изделия)
        окно_продажи.title("Продажа товара")
        окно_продажи.geometry("500x400")  # Увеличиваем размер окна
        окно_продажи.resizable(False, False)

        # Получаем данные о выбранном товаре
        item = self.список_готовых.item(selected_items[0])
        values = item['values']

        # Отображаем информацию о товаре
        frame_информация = ttk.LabelFrame(окно_продажи, text="Информация о товаре", padding=15)  # Увеличиваем отступы
        frame_информация.pack(fill='x', padx=15, pady=10)

        # Увеличиваем размер шрифта для меток
        style = ttk.Style()
        style.configure('Info.TLabel', font=('Arial', 12))

        ttk.Label(frame_информация, text=f"Название: {values[2]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Тип: {values[3]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Размер: {values[4]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Цвет: {values[5]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Доступное количество: {values[6]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Цена: {values[7]}", style='Info.TLabel').pack(anchor='w', pady=3)

        # Поле для ввода количества
        frame_количество = ttk.LabelFrame(окно_продажи, text="Количество для продажи", padding=15)
        frame_количество.pack(fill='x', padx=15, pady=10)

        количество = ttk.Entry(frame_количество, width=15, font=('Arial', 12))  # Увеличиваем размер поля
        количество.pack(pady=10)
        количество.insert(0, "1")

        def подтвердить_продажу():
            try:
                кол = int(количество.get())
                if кол <= 0:
                    raise ValueError("Количество должно быть положительным числом")

                доступное_количество = int(values[6])
                if кол > доступное_количество:
                    raise ValueError(f"Недостаточно товара. Доступно: {доступное_количество}")

                # Получаем ID товара из базы данных
                self.cursor.execute("""
                    SELECT id, price FROM clothing
                    WHERE name = %s AND type = %s AND size = %s AND color = %s AND material = 'Ассортимент магазина'
                    LIMIT 1
                """, (values[2], values[3], values[4], values[5]))
                row = self.cursor.fetchone()
                if not row:
                    raise ValueError("Товар не найден в базе данных")

                clothing_id, price = row
                total_price = price * кол

                # Добавляем запись о продаже
                self.cursor.execute("""
                    INSERT INTO sales (clothing_id, type, size, color, price, quantity, total_price)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (clothing_id, values[3], values[4], values[5], price, кол, total_price))

                # Уменьшаем количество товара
                self.cursor.execute("""
                    UPDATE clothing
                    SET quantity = quantity - %s
                    WHERE id = %s
                """, (кол, clothing_id))

                self.conn.commit()

                # Обновляем списки
                self.обновить_список_готовых_изделий()
                self.обновить_список_продаж()

                messagebox.showinfo("Успех", f"Товар успешно продан!\nКоличество: {кол}\nОбщая сумма: {total_price} руб.")
                окно_продажи.destroy()

            except ValueError as e:
                messagebox.showerror("Ошибка", str(e))
            except Error as e:
                self.conn.rollback()
                messagebox.showerror("Ошибка базы данных", str(e))

        # Кнопки
        frame_кнопки = ttk.Frame(окно_продажи)
        frame_кнопки.pack(pady=15)
        ttk.Button(frame_кнопки, text="Подтвердить продажу", command=подтвердить_продажу, width=20).pack(side='left', padx=10)
        ttk.Button(frame_кнопки, text="Отмена", command=окно_продажи.destroy, width=20).pack(side='left', padx=10)

        # Делаем окно модальным
        окно_продажи.transient(self.frame_готовые_изделия)
        окно_продажи.grab_set()
        self.frame_готовые_изделия.wait_window(окно_продажи)

    def создать_интерфейс_продаж(self):
        # Основной контейнер
        main_frame = ttk.Frame(self.frame_продажи)
        main_frame.pack(expand=True, fill='both', padx=20, pady=20)

        # Список продаж
        frame_список = ttk.LabelFrame(main_frame, text="История продаж", padding=10)
        frame_список.pack(fill='both', expand=True, padx=5, pady=5)

        self.список_продаж = ttk.Treeview(frame_список, columns=(
            "Название", "Тип", "Размер", "Цвет", "Количество", "Цена за ед.", "Общая сумма", "Дата продажи"
        ), show="headings")
        self.список_продаж.pack(expand=True, fill='both')

        # Настраиваем ширину колонок
        self.список_продаж.column("Название", width=200)
        self.список_продаж.column("Тип", width=150)
        self.список_продаж.column("Размер", width=100)
        self.список_продаж.column("Цвет", width=100)
        self.список_продаж.column("Количество", width=100)
        self.список_продаж.column("Цена за ед.", width=100)
        self.список_продаж.column("Общая сумма", width=100)
        self.список_продаж.column("Дата продажи", width=150)

        for col in ("Название", "Тип", "Размер", "Цвет", "Количество", "Цена за ед.", "Общая сумма", "Дата продажи"):
            self.список_продаж.heading(col, text=col)

        # Метка для общей суммы продаж
        self.метка_суммы_продаж = ttk.Label(main_frame, text="", font=('Arial', 12, 'bold'))
        self.метка_суммы_продаж.pack(pady=10)

        # Обновляем список продаж
        self.обновить_список_продаж()

    def обновить_список_продаж(self):
        # Очищаем список
        for item in self.список_продаж.get_children():
            self.список_продаж.delete(item)

        # Получаем данные о продажах
        self.cursor.execute("""
            SELECT c.name, s.type, s.size, s.color, s.quantity, s.price, s.total_price, s.sale_date
            FROM sales s
            JOIN clothing c ON s.clothing_id = c.id
            ORDER BY s.sale_date DESC
        """)

        total_sales = 0
        for продажа in self.cursor.fetchall():
            self.список_продаж.insert("", "end", values=(
                продажа[0],  # Название
                продажа[1],  # Тип
                продажа[2],  # Размер
                продажа[3],  # Цвет
                продажа[4],  # Количество
                f"{продажа[5]} руб.",  # Цена за ед.
                f"{продажа[6]} руб.",  # Общая сумма
                продажа[7].strftime("%Y-%m-%d %H:%M")  # Дата продажи
            ))
            total_sales += продажа[6]

        self.метка_суммы_продаж.configure(text=f"Общая сумма продаж: {total_sales:.2f} руб.")

    def показать_информацию_о_товаре(self, event):
        # Получаем выбранный элемент
        selection = self.список_готовых.selection()
        if not selection:
            return

        # Получаем значения выбранного элемента
        values = self.список_готовых.item(selection[0])['values']

        # Создаем новое окно
        info_window = tk.Toplevel(self.root)
        info_window.title("Информация о товаре")
        info_window.geometry("400x500")
        info_window.resizable(False, False)

        # Создаем фрейм для информации
        frame_информация = ttk.Frame(info_window, padding=20)
        frame_информация.pack(fill='both', expand=True)

        # Стиль для меток
        style = ttk.Style()
        style.configure('Info.TLabel', font=('Arial', 12))

        ttk.Label(frame_информация, text=f"ID: {values[0]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Штрих-код: {values[1]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Название: {values[2]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Тип: {values[3]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Размер: {values[4]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Цвет: {values[5]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Доступное количество: {values[6]}", style='Info.TLabel').pack(anchor='w', pady=3)
        ttk.Label(frame_информация, text=f"Цена: {values[7]}", style='Info.TLabel').pack(anchor='w', pady=3)

    def сбросить_поиск(self):
        self.поиск_штрих_код.delete(0, tk.END)
        self.обновить_список_готовых_изделий()

def все_заполнено(поля):
    return all(поле.strip() for поле in поля)

if __name__ == "__main__":
    root = tk.Tk()
    app = ТрикотажнаяФабрикаGUI(root)
    root.mainloop()
